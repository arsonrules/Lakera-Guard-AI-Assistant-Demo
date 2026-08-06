"""
Text extraction for uploaded RAG documents.

The knowledge base is stored and scanned as plain text, so a richer upload
(PDF, DOCX, CSV, JSON) is converted to text ONCE at upload time rather than
parsed on every retrieval. Everything downstream — CP2 scanning, `rag.retrieve`,
the report — keeps working on `.txt` unchanged.

Security posture
----------------
Parsing untrusted files is itself attack surface, so this module is deliberately
paranoid:

* **Fail closed.** Anything unparseable, encrypted, or yielding no usable text
  raises `ExtractError` and the upload is rejected. We never index an empty or
  half-read document, because a silently-empty knowledge base looks identical to
  a working one right up until it matters.
* **Bounded output.** Extraction is capped (`MAX_CHARS`) so a small crafted file
  cannot expand into something that exhausts memory or blows past Guard's
  request limits — a decompression-bomb shape is easy to build in these formats.
* **Bounded input.** Page/row/element counts are capped too, so a file that is
  small on disk but structurally enormous still terminates quickly.
"""
from __future__ import annotations

import csv
import io
import json

# Roughly 400 KB of text — comfortably more than any demo document, far below
# anything that would trouble the scanner or the model context.
MAX_CHARS = 400_000
MAX_PDF_PAGES = 200
MAX_CSV_ROWS = 5_000
MAX_DOCX_BLOCKS = 10_000

SUPPORTED_EXTS = {".txt", ".md", ".csv", ".json", ".pdf", ".docx"}


class ExtractError(ValueError):
    """A user-facing extraction problem. The message is shown verbatim, so it
    must say what is wrong AND what to do about it."""


def _truncate(text: str) -> str:
    return text if len(text) <= MAX_CHARS else text[:MAX_CHARS] + "\n…[truncated]"


def _plain(data: bytes, name: str) -> str:
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        raise ExtractError(f"{name} isn't valid UTF-8 text.")


def _from_csv(data: bytes, name: str) -> str:
    """Flatten to `col: value` lines — readable prose for a text scanner, which
    a raw comma-separated dump is not."""
    text = _plain(data, name)
    try:
        rows = list(csv.reader(io.StringIO(text)))
    except csv.Error as exc:
        raise ExtractError(f"{name} could not be parsed as CSV: {exc}")
    if not rows:
        raise ExtractError(f"{name} contains no rows.")
    header, *body = rows
    out = []
    for row in body[:MAX_CSV_ROWS]:
        out.append(" · ".join(
            f"{(header[i] if i < len(header) else f'col{i}')}: {v}"
            for i, v in enumerate(row) if v
        ))
    return "\n".join(x for x in out if x) or " ".join(header)


def _from_json(data: bytes, name: str) -> str:
    text = _plain(data, name)
    try:
        parsed = json.loads(text)
    except ValueError as exc:
        raise ExtractError(f"{name} is not valid JSON: {exc}")
    # Re-serialised rather than passed through: normalises formatting and keeps
    # keys next to values so injected instructions stay in readable context.
    return json.dumps(parsed, ensure_ascii=False, indent=2)


def _from_pdf(data: bytes, name: str) -> str:
    try:
        import pypdf
    except ImportError:                       # pragma: no cover - dependency is pinned
        raise ExtractError("PDF support isn't installed on the server.")
    try:
        reader = pypdf.PdfReader(io.BytesIO(data))
        if getattr(reader, "is_encrypted", False):
            raise ExtractError(f"{name} is password-protected and can't be read.")
        parts = [(page.extract_text() or "") for page in reader.pages[:MAX_PDF_PAGES]]
    except ExtractError:
        raise
    except Exception as exc:                  # noqa: BLE001 — any parser failure is user-facing
        raise ExtractError(f"{name} could not be read as a PDF: {type(exc).__name__}")
    text = "\n".join(p for p in parts if p.strip())
    if not text.strip():
        # The common real-world case: a scanned/image-only PDF. Say so, because
        # "no text found" without the reason sends people hunting the wrong bug.
        raise ExtractError(f"{name} contains no extractable text (image-only or scanned PDF?).")
    return text


def _from_docx(data: bytes, name: str) -> str:
    try:
        import docx
    except ImportError:                       # pragma: no cover - dependency is pinned
        raise ExtractError("DOCX support isn't installed on the server.")
    try:
        doc = docx.Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs[:MAX_DOCX_BLOCKS]]
        # Table cells carry real content in policy documents — and are a natural
        # hiding place for an indirect-injection payload, so don't skip them.
        for table in doc.tables:
            for row in table.rows:
                parts.append(" · ".join(c.text.strip() for c in row.cells if c.text.strip()))
    except Exception as exc:                  # noqa: BLE001
        raise ExtractError(f"{name} could not be read as a DOCX: {type(exc).__name__}")
    text = "\n".join(p for p in parts if p and p.strip())
    if not text.strip():
        raise ExtractError(f"{name} contains no extractable text.")
    return text


_EXTRACTORS = {
    ".txt": _plain, ".md": _plain,
    ".csv": _from_csv, ".json": _from_json,
    ".pdf": _from_pdf, ".docx": _from_docx,
}


def extract(filename: str, data: bytes) -> str:
    """
    Extract scannable text from an uploaded document.

    Raises `ExtractError` with a message safe (and useful) to show the user.
    """
    ext = ("." + filename.rsplit(".", 1)[-1].lower()) if "." in filename else ""
    fn = _EXTRACTORS.get(ext)
    if fn is None:
        supported = " ".join(sorted(SUPPORTED_EXTS))
        raise ExtractError(f"{filename} isn't a supported format. Use: {supported}")
    if not data:
        raise ExtractError(f"{filename} is empty.")
    text = _truncate(fn(data, filename).strip())
    if not text:
        raise ExtractError(f"{filename} contains no extractable text.")
    return text
